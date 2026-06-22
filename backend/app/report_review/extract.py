"""把上传的 Word / Excel / Markdown 文件解析为带稳定锚点的规范化结构.

每个文件 -> ReportDoc，包含一串 blocks。
- docx: 每个非空段落 / 表格 = 一个 block，锚点 = "p{i}" / "t{i}"
- xlsx: 每个 sheet 的非空行抽样 = block，锚点 = "{sheet}!r{row}"
- md/txt: 按行/段，锚点 = "L{i}"

前端用 block.anchor 渲染原文并支持「点击复核意见 -> 高亮定位」。
LLM 复核时把 (file_id, anchor, text) 一并喂进去，让它在 source_refs 里引用锚点。
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field, asdict
from typing import Any


# ── 支持的扩展名 ──────────────────────────────────────────────
DOCX_EXT = (".docx",)
XLSX_EXT = (".xlsx", ".xlsm")
TEXT_EXT = (".md", ".markdown", ".txt")

# 去除会破坏 JSON.parse 的控制字符（保留 \t）
_CTRL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _clean(text: str) -> str:
    """清洗从文档抽取的文本：去控制字符、归一空白。"""
    if not text:
        return ""
    # \x0b / \x0c（Word 软换行/分页）转空格，其余控制字符删除
    text = text.replace("\x0b", " ").replace("\x0c", " ")
    text = _CTRL_RE.sub("", text)
    return text.strip()


@dataclass
class Block:
    anchor: str                 # 文件内唯一锚点, e.g. "p7" / "审定报表!r12"
    kind: str                   # paragraph | heading | table | sheet_row | line
    text: str                   # 纯文本（表格被展平为多行）
    meta: dict[str, Any] = field(default_factory=dict)   # 额外信息（sheet 名、行号、表格 rows…）


@dataclass
class ReportDoc:
    file_id: str                # 稳定 id（前端用来定位）
    filename: str
    kind: str                   # word | excel | markdown | unknown
    blocks: list[Block] = field(default_factory=list)
    note: str = ""              # 解析备注 / 错误

    def to_dict(self) -> dict[str, Any]:
        return {
            "file_id": self.file_id,
            "filename": self.filename,
            "kind": self.kind,
            "note": self.note,
            "blocks": [asdict(b) for b in self.blocks],
        }

    def plain_text(self, max_chars: int = 12000) -> str:
        """把全部 block 拼成带锚点标注的纯文本，喂给 LLM。"""
        out: list[str] = []
        used = 0
        for b in self.blocks:
            line = f"[{b.anchor}] {b.text}"
            if used + len(line) > max_chars:
                out.append("…（原文过长已截断）")
                break
            out.append(line)
            used += len(line)
        return "\n".join(out)


def _ext(filename: str) -> str:
    name = (filename or "").lower()
    dot = name.rfind(".")
    return name[dot:] if dot >= 0 else ""


# ── Word ──────────────────────────────────────────────────────
def _parse_docx(file_id: str, filename: str, raw: bytes) -> ReportDoc:
    doc = ReportDoc(file_id=file_id, filename=filename, kind="word")
    try:
        from docx import Document  # python-docx
    except ImportError:
        doc.note = "未安装 python-docx，无法解析 Word"
        return doc
    try:
        d = Document(io.BytesIO(raw))
    except Exception as e:  # noqa: BLE001
        doc.note = f"Word 解析失败: {e}"
        return doc

    # 段落
    for i, p in enumerate(d.paragraphs):
        t = _clean(p.text or "")
        if not t:
            continue
        style = (p.style.name if p.style else "") or ""
        is_heading = style.lower().startswith("heading") or style.startswith("标题")
        doc.blocks.append(Block(
            anchor=f"p{i}",
            kind="heading" if is_heading else "paragraph",
            text=t,
            meta={"style": style},
        ))

    # 表格（展平为 "行: a | b | c"）
    for ti, table in enumerate(d.tables):
        rows_text: list[str] = []
        rows_data: list[list[str]] = []
        for r in table.rows:
            cells = [_clean((c.text or "").replace("\n", " ")) for c in r.cells]
            if not any(cells):
                continue
            rows_data.append(cells)
            rows_text.append(" | ".join(cells))
        if not rows_text:
            continue
        # 表格做一个汇总 block（便于 LLM 把整张表当上下文）
        preview = "\n".join(rows_text[:40])
        doc.blocks.append(Block(
            anchor=f"t{ti}",
            kind="table",
            text=f"【表格{ti + 1}】\n{preview}",
            meta={"rows": rows_data[:200], "row_count": len(rows_data)},
        ))

    if not doc.blocks:
        doc.note = "Word 内未提取到文本"
    return doc


# ── Excel ─────────────────────────────────────────────────────
def _cell_str(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return _clean(str(v))


def _parse_xlsx(file_id: str, filename: str, raw: bytes) -> ReportDoc:
    doc = ReportDoc(file_id=file_id, filename=filename, kind="excel")
    try:
        import openpyxl
    except ImportError:
        doc.note = "未安装 openpyxl，无法解析 Excel"
        return doc
    try:
        wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    except Exception as e:  # noqa: BLE001
        doc.note = f"Excel 解析失败: {e}"
        return doc

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        row_idx = 0
        emitted = 0
        for row in ws.iter_rows(values_only=True):
            row_idx += 1
            cells = [_cell_str(c) for c in row]
            # 只保留有内容的行，避免噪音
            non_empty = [c for c in cells if c]
            if not non_empty:
                continue
            # 截掉尾部空列
            while cells and cells[-1] == "":
                cells.pop()
            text = " | ".join(cells)
            if not text.strip():
                continue
            doc.blocks.append(Block(
                anchor=f"{sheet_name}!r{row_idx}",
                kind="sheet_row",
                text=text,
                meta={"sheet": sheet_name, "row": row_idx},
            ))
            emitted += 1
            if emitted >= 400:        # 每个 sheet 最多 400 行，防止超大表炸内存
                doc.blocks.append(Block(
                    anchor=f"{sheet_name}!rMORE",
                    kind="sheet_row",
                    text=f"…（{sheet_name} 行数过多，已截断）",
                    meta={"sheet": sheet_name, "truncated": True},
                ))
                break
    wb.close()
    if not doc.blocks:
        doc.note = "Excel 内未提取到数据"
    return doc


# ── Markdown / 文本 ───────────────────────────────────────────
def _parse_text(file_id: str, filename: str, raw: bytes) -> ReportDoc:
    doc = ReportDoc(file_id=file_id, filename=filename, kind="markdown")
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = raw.decode("gbk")
        except Exception:  # noqa: BLE001
            text = raw.decode("utf-8", errors="ignore")
    for i, line in enumerate(text.splitlines()):
        t = _clean(line)
        if not t:
            continue
        is_heading = t.lstrip().startswith("#")
        doc.blocks.append(Block(
            anchor=f"L{i + 1}",
            kind="heading" if is_heading else "line",
            text=t,
            meta={"line": i + 1},
        ))
    if not doc.blocks:
        doc.note = "文本为空"
    return doc


def parse_upload(file_id: str, filename: str, raw: bytes) -> ReportDoc:
    """根据扩展名分派解析。"""
    ext = _ext(filename)
    if ext in DOCX_EXT:
        return _parse_docx(file_id, filename, raw)
    if ext in XLSX_EXT:
        return _parse_xlsx(file_id, filename, raw)
    if ext in TEXT_EXT:
        return _parse_text(file_id, filename, raw)
    # 老 .doc / .xls 不支持
    doc = ReportDoc(file_id=file_id, filename=filename, kind="unknown")
    doc.note = f"暂不支持的文件类型 {ext or '(无扩展名)'}（支持 .docx/.xlsx/.xlsm/.md/.txt）"
    return doc


def read_instruction_file(filename: str, raw: bytes) -> str:
    """复核要求文档（.md/.txt/.docx）-> 纯文本。"""
    ext = _ext(filename)
    if ext in DOCX_EXT:
        d = _parse_docx("instr", filename, raw)
        return "\n".join(b.text for b in d.blocks)
    # 当作文本
    d = _parse_text("instr", filename, raw)
    return "\n".join(b.text for b in d.blocks)
